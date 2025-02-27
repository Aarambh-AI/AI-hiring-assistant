from docx import Document
import fitz
import io
import pandas as pd

def convert_pdf_to_text(pdf_data):
    # Open the PDF file using PyMuPDF from a byte stream
    pdf_stream = io.BytesIO(pdf_data)
    document = fitz.open(stream=pdf_stream, filetype="pdf")
    text = ""
    
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += page.get_text()
    
    return text

def convert_docx_to_text(docx_data):
    # Open the DOCX file using python-docx from a byte stream
    docx_stream = io.BytesIO(docx_data)
    document = Document(docx_stream)
    text = ""
    # Extract text and table content
    full_text = []
    
    # Process paragraphs
    for para in document.paragraphs:
        full_text.append(para.text)
    
    # Process tables
    for table in document.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            table_text.append(' | '.join(row_text))
        full_text.extend(table_text)
    
    return '\n'.join(full_text)

def convert_text_file_to_text(text_data):
    # Directly decode the text file data
    return text_data.decode('utf-8')

def convert_csv_to_text(csv_data):
    # Open the CSV file using pandas from a byte stream
    csv_stream = io.BytesIO(csv_data)
    
    try:
        # Try to read with default settings first
        df = pd.read_csv(csv_stream)
    except UnicodeDecodeError:
        # If default encoding fails, try with different encodings
        csv_stream.seek(0)  # Reset stream position
        df = pd.read_csv(csv_stream, encoding='latin1')
    
    # Convert DataFrame to list of dictionaries
    # Replace NaN values with None and convert all other values to strings
    records = df.replace({pd.NA: None}).astype(str).to_dict('records')
    
    # Convert the list of dictionaries to a string representation
    return records

def convert_excel_to_text(excel_data):
    # Open the Excel file using pandas from a byte stream
    excel_stream = io.BytesIO(excel_data)
    
    # Read all sheets from the Excel file
    dfs = pd.read_excel(excel_stream, sheet_name=None)
    
    # Initialize list to store results from all sheets
    all_records = []
    
    # Process each sheet
    for sheet_name, df in dfs.items():
        # Convert DataFrame to list of dictionaries
        # Replace NaN values with None and convert all other values to strings
        records = df.replace({pd.NA: None}).astype(str).to_dict('records')
        
        # Add sheet name to each record
        for record in records:
            record['sheet_name'] = sheet_name
        
        all_records.extend(records)
    
    return all_records

def convert_file_to_text(file_name, file_data):
    # Determine file type based on extension and convert accordingly
    if file_name.lower().endswith('.pdf'):
        return convert_pdf_to_text(file_data)
    elif file_name.lower().endswith('.docx'):
        return convert_docx_to_text(file_data)
    elif file_name.lower().endswith('.txt'):
        return convert_text_file_to_text(file_data)
    elif file_name.lower().endswith(('.xlsx', '.xls')):
        return convert_excel_to_text(file_data)
    elif file_name.lower().endswith('.csv'):
        return convert_csv_to_text(file_data)
    else:
        raise ValueError("Unsupported file type")
    
def main():
    # Test files
    test_files = [
        ('sample.txt', 'Hello World!'.encode('utf-8')),
        ('sample.pdf', open('path/to/sample.pdf', 'rb').read()),
        ('sample.docx', open('path/to/sample.docx', 'rb').read())
    ]

    # Test each file conversion
    for file_name, file_data in test_files:
        try:
            print(f"\nProcessing {file_name}...")
            text = convert_file_to_text(file_name, file_data)
            print("Converted text:")
            print("-" * 50)
            print(text)
            print("-" * 50)
        except Exception as e:
            print(f"Error processing {file_name}: {str(e)}")

if __name__ == "__main__":
    main()
    

